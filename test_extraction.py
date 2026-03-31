import io
from docx import Document
import pandas as pd
import fitz
import os

def test_extraction():
    # 1. Test Text
    text_content = "This is a test text file about Google."
    file_txt = io.BytesIO(text_content.encode('utf-8'))
    file_txt.filename = "test.txt"
    
    # 2. Test CSV
    df = pd.DataFrame({'name': ['Varun', 'DeepMind'], 'city': ['Jaipur', 'London']})
    file_csv = io.BytesIO()
    df.to_csv(file_csv, index=False)
    file_csv.seek(0)
    file_csv.filename = "test.csv"
    
    # 3. Test DOCX
    doc = Document()
    doc.add_paragraph("Hello from a Word document. Microsoft is a company.")
    file_docx = io.BytesIO()
    doc.save(file_docx)
    file_docx.seek(0)
    file_docx.filename = "test.docx"
    
    # 4. Test PDF
    pdf_doc = fitz.open()
    page = pdf_doc.new_page()
    page.insert_text((50, 50), "This is a PDF document. Apple is in California.")
    pdf_bytes = pdf_doc.write()
    file_pdf = io.BytesIO(pdf_bytes)
    file_pdf.filename = "test.pdf"
    
    # Mocking the extract_text function logic
    def mock_extract(file):
        filename = file.filename.lower()
        file_bytes = file.read()
        if filename.endswith('.pdf'):
            text = ""
            with fitz.open(stream=file_bytes, filetype="pdf") as doc:
                for page in doc: text += page.get_text()
            return text
        elif filename.endswith('.docx'):
            d = Document(io.BytesIO(file_bytes))
            return "\n".join([p.text for p in d.paragraphs])
        elif filename.endswith('.csv'):
            return pd.read_csv(io.BytesIO(file_bytes)).to_string(index=False)
        else:
            return file_bytes.decode('utf-8', errors='ignore')

    print("--- Testing TXT ---")
    print(mock_extract(file_txt))
    print("\n--- Testing CSV ---")
    print(mock_extract(file_csv))
    print("\n--- Testing DOCX ---")
    print(mock_extract(file_docx))
    print("\n--- Testing PDF ---")
    print(mock_extract(file_pdf))

if __name__ == "__main__":
    test_extraction()
